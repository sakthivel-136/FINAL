import streamlit as st
import base64
import os
import pandas as pd
import pickle
import smtplib
import re
from email.message import EmailMessage
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from gtts import gTTS
import tempfile
from deep_translator import GoogleTranslator
import time
from docx import Document
from docx.shared import Pt, RGBColor
from PIL import Image
import chatbot_main

# ========== PAGE 1 ==========
if "page" not in st.session_state:
    st.session_state.page = 1
if "img_idx" not in st.session_state:
    st.session_state.img_idx = 0
if "autoplay_enabled" not in st.session_state:
    st.session_state.autoplay_enabled = True

if st.session_state.page == 1:
    st.set_page_config(page_title="KCET Welcome", layout="centered")

    # Load Logo
    try:
        with open("kcet_logo.png", "rb") as img_file:
            logo_base64 = base64.b64encode(img_file.read()).decode()
    except:
        logo_base64 = ""

    st.markdown("""
        <div style='text-align:center; padding: 30px;'>
            <img src='data:image/png;base64,""" + logo_base64 + """' style='height:100px; width:100px; border-radius:50%;'>
            <h2 style='margin-top:10px;'>KAMARAJ COLLEGE OF ENGINEERING AND TECHNOLOGY</h2>
        </div>
    """, unsafe_allow_html=True)

    # Add background music
    if os.path.exists("kcet_music.mp3"):
        audio_file = open("kcet_music.mp3", "rb")
        audio_bytes = audio_file.read()
        audio_base64 = base64.b64encode(audio_bytes).decode("utf-8")
        st.markdown(f"""
            <audio autoplay loop>
                <source src="data:audio/mp3;base64,{audio_base64}" type="audio/mp3">
            </audio>
        """, unsafe_allow_html=True)

    # Slideshow with autoplay every 3 seconds
    image_folder = "college_images"
    images = [f for f in os.listdir(image_folder) if f.lower().endswith((".png", ".jpg", ".jpeg"))]

    captions = [os.path.splitext(img)[0].replace("_", " ") for img in images]

    if images:
        slideshow = st.empty()
        caption_holder = st.empty()

        for _ in range(1):  # Run only once unless rerun
            current_index = st.session_state.img_idx
            image_path = os.path.join(image_folder, images[current_index])
            caption = captions[current_index]
            try:
                img = Image.open(image_path)
                target_ratio = 16 / 9
                width, height = img.size
                current_ratio = width / height
                if current_ratio > target_ratio:
                    new_width = int(target_ratio * height)
                    offset = (width - new_width) // 2
                    img = img.crop((offset, 0, offset + new_width, height))
                elif current_ratio < target_ratio:
                    new_height = int(width / target_ratio)
                    offset = (height - new_height) // 2
                    img = img.crop((0, offset, width, offset + new_height))

                with slideshow.container():
                    st.markdown("<div style='animation:fadein 1s;'>", unsafe_allow_html=True)
                    st.image(img, use_container_width=True)
                    st.markdown("</div>", unsafe_allow_html=True)
                with caption_holder.container():
                    st.markdown(f"<div style='text-align:center; font-size:20px; margin-top:10px; color:#333;'>{caption}</div>", unsafe_allow_html=True)

            except:
                st.warning(f"Could not load image: {image_path}")

            time.sleep(3)
            st.session_state.img_idx = (current_index + 1) % len(images)

    # Go to chatbot button
    if st.button("Go to Chatbot", help="Enter the assistant page"):
        st.session_state.page = 99
        st.session_state.autoplay_enabled = False
        st.rerun()

# ========== PAGE 99 (Loader before Chatbot) ==========
elif st.session_state.page == 99:
    st.set_page_config(page_title="Loading KCET Chatbot", layout="centered")
    st.markdown("""
        <div style='text-align:center; margin-top:100px;'>
            <img src='https://i.gifer.com/VAyR.gif' width='100'>
            <h4 style='margin-top:20px;'>Launching KCET Chatbot...</h4>
        </div>
    """, unsafe_allow_html=True)
    time.sleep(2.5)
    st.session_state.page = 2
    st.rerun()

# ========== PAGE 2 (CHATBOT) ==========
elif st.session_state.page == 2:
    st.set_page_config(page_title="KCET Chatbot", layout="centered")

    # Button to return to main page
    if st.button("\U0001f3e0 Main Page"):
        st.session_state.page = 1
        st.rerun()

    st.markdown("""
        <div style='text-align:center; margin-top: 50px;'>
            <h3>🤖 KCET Chatbot is now active!</h3>
        </div>
    """, unsafe_allow_html=True)

    # Run chatbot logic
    chatbot_main.run_chatbot()



    # ========== EMAIL CREDENTIALS ==========
    SENDER_EMAIL = "kamarajengg.edu.in@gmail.com"
    SENDER_PASSWORD = "vwvcwsfffbrvumzh"  # Gmail App Password

    tf_vector_file = "vectorized.pkl"
    csv_file = "kcet.csv"
    threshold = 0.6

    def remove_emojis(text):
        return re.sub(r'[^\x00-\x7F]+', '', text)

    def speak_text(text):
        lang = 'ta' if st.session_state.get("language", "en") == "ta" else 'en'
        tts = gTTS(text=text, lang=lang, tld='com')
        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as fp:
            tts.save(fp.name)
            audio_base64 = base64.b64encode(open(fp.name, "rb").read()).decode()
            audio_html = f"""
                <audio autoplay>
                    <source src="data:audio/mp3;base64,{audio_base64}" type="audio/mp3">
                </audio>
            """
            st.markdown(audio_html, unsafe_allow_html=True)

    def send_email(recipient_emails, subject, body, attachment_path):
        if isinstance(recipient_emails, str):
            recipient_emails = [email.strip() for email in recipient_emails.split(",") if "@" in email]
        msg = EmailMessage()
        msg["Subject"] = subject
        msg["From"] = SENDER_EMAIL
        msg["To"] = ", ".join(recipient_emails)
        msg.set_content(body)
        with open(attachment_path, "rb") as f:
            msg.add_attachment(f.read(), maintype="application", subtype="octet-stream", filename=os.path.basename(attachment_path))
        try:
            with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
                smtp.login(SENDER_EMAIL, SENDER_PASSWORD)
                smtp.send_message(msg)
            return True
        except Exception as e:
            return str(e)

    @st.cache_data
    def load_vector_data():
        try:
            if os.path.exists(tf_vector_file):
                with open(tf_vector_file, "rb") as f:
                    vectorizer, vectors, df = pickle.load(f)
            else:
                raise Exception("File missing or invalid format")
        except:
            df = pd.read_csv(csv_file)
            df['Question'] = df['Question'].str.lower().str.strip()
            vectorizer = TfidfVectorizer()
            vectors = vectorizer.fit_transform(df['Question'])
            with open(tf_vector_file, "wb") as f:
                pickle.dump((vectorizer, vectors, df), f)
        return vectorizer, vectors, df

    vectorizer, vectors, df = load_vector_data()

    for key, default in [
        ("language", "en"),
        ("original_log", []),
        ("last_input", ""),
        ("username", "You"),
        ("user_color", "#d0e8f2"),
        ("bot_color", "#d1d1e9"),
        ("enable_export", True),
    ]:
        if key not in st.session_state:
            st.session_state[key] = default

    st.sidebar.header("⚙️ Chat Settings")
    st.session_state.username = st.sidebar.text_input("🧑 Your Name", value=st.session_state.username)
    st.session_state.user_color = st.sidebar.color_picker("🎨 User Bubble Color", value=st.session_state.user_color)
    st.session_state.bot_color = st.sidebar.color_picker("🎨 Bot Bubble Color", value=st.session_state.bot_color)
    st.session_state.enable_export = st.sidebar.checkbox("📤 Enable Chat Export", value=st.session_state.enable_export)

    try:
        with open("kcet_logo.png", "rb") as img_file:
            logo_base64 = base64.b64encode(img_file.read()).decode()
    except:
        logo_base64 = ""

    st.markdown(f"""
    <div style='display:flex; justify-content:space-between; align-items:center; background:#333; padding:8px 16px; border-radius:6px;'>
        <div style='display:flex; align-items:center;'>
            <img src='data:image/png;base64,{logo_base64}' style='height:40px; width:40px; border-radius:50%; margin-right:10px;'>
            <span style='color:white; font-size:14px; font-weight:bold;'>KAMARAJ COLLEGE OF ENGINEERING AND TECHNOLOGY</span>
        </div>
        <marquee scrollamount='5' style='color:white; font-size:13px;'>💼 100% Placement | 👩‍🏫 Top Faculty | 🎓 Research Driven | 🧠 Hackathons | 🤝 Industry Collaboration</marquee>
    </div>
    """, unsafe_allow_html=True)

    with st.form("chat_form", clear_on_submit=True):
        user_input = st.text_input("Ask your question...", value=st.session_state.last_input)
        submitted = st.form_submit_button("➔")
        if submitted:
            st.session_state.last_input = ""

    if submitted and user_input.strip():
        msg = remove_emojis(user_input.strip())
        st.session_state.original_log.append((st.session_state.username, msg, "User"))
        vec = vectorizer.transform([msg.lower()])
        similarity = cosine_similarity(vec, vectors)
        idx = similarity.argmax()
        max_sim = similarity.max()
        answer = df.iloc[idx]['Answer'] if max_sim >= threshold else "❌ Sorry, I couldn't understand that. Please rephrase."
        answer = remove_emojis(answer)
        if st.session_state.language == "ta":
            answer = GoogleTranslator(source='en', target='ta').translate(answer)
        st.session_state.original_log.append(("KCET Assistant", answer, "Assistant"))
        with st.spinner("KCET Assistant typing..."):
            time.sleep(min(1.5, len(answer)/20))
            speak_text(answer)

    def translate_all():
        for i in range(len(st.session_state.original_log)):
            speaker, msg, role = st.session_state.original_log[i]
            if role == "Assistant":
                translated = GoogleTranslator(source='en' if st.session_state.language=='ta' else 'ta', target=st.session_state.language).translate(msg)
                st.session_state.original_log[i] = (speaker, translated, role)
        st.rerun()

    def clear_chat():
        st.session_state.original_log.clear()
        st.session_state.last_input = ""
        st.rerun()

    for speaker, msg, role in st.session_state.original_log:
        align = 'right' if role == "User" else 'left'
        color = st.session_state.user_color if role == "User" else st.session_state.bot_color
        st.markdown(f"""
            <div style='background-color:{color}; padding:10px; margin:10px; border-radius:10px; text-align:{align};'>
                <b>{speaker}</b>: {msg}
            </div>
        """, unsafe_allow_html=True)

    if st.session_state.enable_export:
        with st.expander("📤 Export Chat as TXT/DOC and Email"):
            file_format = st.radio("Select Format", ["TXT", "DOC"])
            recipients = st.text_input("📧 Enter comma-separated emails")

            def export_chat(ext):
                content = [(s, r, m) for s, m, r in st.session_state.original_log]
                if ext == "DOC":
                    doc = Document()
                    doc.add_heading("KCET Chat Export", 0)
                    style = doc.styles['Normal']
                    font = style.font
                    font.name = 'Calibri'
                    font.size = Pt(11)
                    font.color.rgb = RGBColor(0, 0, 0)
                    for speaker, role, message in content:
                        p = doc.add_paragraph()
                        p.add_run(f"{speaker} ({role}): ").bold = True
                        p.add_run(message)
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                        doc.save(tmp.name)
                        return tmp.name
                else:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".txt", mode="w", encoding="utf-8") as tmp:
                        tmp.write("\n".join([f"{s} ({r}): {m}" for s, r, m in content]))
                        return tmp.name

            if st.button("Send Email"):
                if recipients and file_format:
                    file_path = export_chat(file_format)
                    result = send_email(
                        recipients,
                        "KCET Chat Export",
                        f"Attached is your KCET chat in {file_format} format.",
                        file_path
                    )
                    if result is True:
                        st.success("✅ Email sent successfully")
                    else:
                        st.error(f"❌ Error: {result}")
                else:
                    st.warning("⚠️ Please enter valid emails and select a format.")

    col1, col2 = st.columns(2)
    with col1:
        if st.button("🔄 Translate to Tamil" if st.session_state.language == "en" else "🖙 Back to English"):
            st.session_state.language = "ta" if st.session_state.language == "en" else "en"
            translate_all()

    with col2:
        if st.button("🩹 Clear Chat"):
            clear_chat()
